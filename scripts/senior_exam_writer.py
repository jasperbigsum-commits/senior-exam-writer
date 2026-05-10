#!/usr/bin/env python3
"""Local evidence-gated senior exam writer.

This script implements a small, inspectable hierarchical RAG workflow:

- SQLite evidence database with FTS5 keyword retrieval.
- Optional llama.cpp local embeddings.
- TOC/overview routing, content retrieval, parent expansion.
- Local llama.cpp generation and strict evidence verification.

It intentionally refuses to generate when evidence is weak.
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import math
import os
import re
import sqlite3
import sys
import textwrap
import time
import urllib.error
import urllib.request
import uuid
import zipfile
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET


SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".docx", ".pdf", ".epub", ".json", ".jsonl"}
DEFAULT_EMBED_URL = os.environ.get("LLAMA_CPP_EMBED_URL", "http://127.0.0.1:8081")
DEFAULT_LLM_URL = os.environ.get("LLAMA_CPP_LLM_URL", "http://127.0.0.1:8080")
DEFAULT_EMBED_MODEL = os.environ.get("LLAMA_CPP_EMBED_MODEL", "local-embedding")
DEFAULT_LLM_MODEL = os.environ.get("LLAMA_CPP_LLM_MODEL", "local-instruct")


@dataclass
class Section:
    title: str
    path: str
    text: str
    locator: str
    level: int = 1


@dataclass
class Evidence:
    id: str
    chunk_id: str
    source_id: str
    layer: str
    path: str
    title: str
    locator: str
    text: str
    score: float
    source_kind: str
    source_title: str
    source_path: str
    source_name: str | None
    source_url: str | None
    published_at: str | None

    def citation(self) -> str:
        source = self.source_title or Path(self.source_path).name
        path = f" - {self.path}" if self.path else ""
        locator = f" - {self.locator}" if self.locator else ""
        if self.source_url or self.published_at or self.source_name:
            parts = [p for p in [self.source_name or source, self.published_at, self.source_url or self.source_path, self.locator] if p]
            return " - ".join(parts)
        return f"{source}{path}{locator}"


def now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat()


def stable_id(*parts: str) -> str:
    h = hashlib.sha256()
    for part in parts:
        h.update(part.encode("utf-8", errors="ignore"))
        h.update(b"\0")
    return h.hexdigest()[:24]


def connect(db_path: str) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("PRAGMA journal_mode = WAL")
    conn.execute("PRAGMA synchronous = NORMAL")
    return conn


def init_db(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS sources (
          id TEXT PRIMARY KEY,
          kind TEXT NOT NULL,
          title TEXT NOT NULL,
          path TEXT,
          source_name TEXT,
          url TEXT,
          published_at TEXT,
          version TEXT,
          metadata_json TEXT NOT NULL DEFAULT '{}',
          created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS chunks (
          id TEXT PRIMARY KEY,
          source_id TEXT NOT NULL REFERENCES sources(id) ON DELETE CASCADE,
          parent_id TEXT REFERENCES chunks(id) ON DELETE CASCADE,
          layer TEXT NOT NULL CHECK(layer IN ('overview', 'toc', 'parent', 'content')),
          path TEXT,
          title TEXT,
          locator TEXT,
          text TEXT NOT NULL,
          token_count INTEGER NOT NULL DEFAULT 0,
          embedding_json TEXT,
          metadata_json TEXT NOT NULL DEFAULT '{}',
          created_at TEXT NOT NULL
        );

        CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(
          chunk_id UNINDEXED,
          title,
          path,
          text,
          tokenize = 'unicode61'
        );

        CREATE TABLE IF NOT EXISTS questions (
          id TEXT PRIMARY KEY,
          topic TEXT NOT NULL,
          question_type TEXT NOT NULL,
          prompt_json TEXT NOT NULL,
          evidence_json TEXT NOT NULL,
          output_json TEXT NOT NULL,
          verification_json TEXT NOT NULL,
          status TEXT NOT NULL,
          created_at TEXT NOT NULL
        );

        CREATE INDEX IF NOT EXISTS idx_chunks_source ON chunks(source_id);
        CREATE INDEX IF NOT EXISTS idx_chunks_parent ON chunks(parent_id);
        CREATE INDEX IF NOT EXISTS idx_chunks_layer ON chunks(layer);
        CREATE INDEX IF NOT EXISTS idx_chunks_path ON chunks(path);
        """
    )
    conn.commit()


def ensure_db(conn: sqlite3.Connection) -> None:
    init_db(conn)


def http_json(url: str, payload: dict[str, Any], timeout: int = 120) -> dict[str, Any]:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            raw = resp.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        details = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"HTTP {exc.code} from {url}: {details}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Cannot reach {url}: {exc}") from exc
    return json.loads(raw)


def llama_embed(texts: list[str], base_url: str, model: str = DEFAULT_EMBED_MODEL) -> list[list[float]]:
    base = base_url.rstrip("/")
    payload = {"model": model, "input": texts}
    try:
        data = http_json(f"{base}/v1/embeddings", payload, timeout=180)
        embeddings = data.get("data", [])
        return [list(map(float, item["embedding"])) for item in embeddings]
    except Exception as first_error:
        if len(texts) != 1:
            vectors = []
            for text in texts:
                vectors.extend(llama_embed([text], base_url, model))
            return vectors
        fallback_payload = {"content": texts[0]}
        try:
            data = http_json(f"{base}/embedding", fallback_payload, timeout=180)
        except Exception as second_error:
            raise RuntimeError(f"embedding failed: {first_error}; fallback failed: {second_error}") from second_error
        emb = data.get("embedding") or data.get("data", [{}])[0].get("embedding")
        if not emb:
            raise RuntimeError(f"embedding response did not include a vector: {data}")
        return [list(map(float, emb))]


def llama_chat(
    messages: list[dict[str, str]],
    base_url: str,
    model: str = DEFAULT_LLM_MODEL,
    temperature: float = 0.2,
    max_tokens: int = 4096,
) -> str:
    base = base_url.rstrip("/")
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False,
    }
    try:
        data = http_json(f"{base}/v1/chat/completions", payload, timeout=300)
        return data["choices"][0]["message"]["content"]
    except Exception as first_error:
        prompt = "\n\n".join(f"{m['role'].upper()}:\n{m['content']}" for m in messages)
        fallback_payload = {
            "prompt": prompt,
            "temperature": temperature,
            "n_predict": max_tokens,
            "stream": False,
        }
        try:
            data = http_json(f"{base}/completion", fallback_payload, timeout=300)
        except Exception as second_error:
            raise RuntimeError(f"generation failed: {first_error}; fallback failed: {second_error}") from second_error
        return data.get("content", "")


def approx_tokens(text: str) -> int:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    words = len(re.findall(r"[A-Za-z0-9_]+", text))
    punctuation = len(re.findall(r"[^\w\s\u4e00-\u9fff]", text))
    return cjk + words + max(1, punctuation // 3)


def normalize_ws(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def split_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"\n{1,}", normalize_ws(text)) if p.strip()]


def read_text_file(path: Path) -> list[tuple[str, str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [("text", text)]


def read_docx(path: Path) -> list[tuple[str, str]]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required to ingest .docx files") from exc
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            style = (para.style.name if para.style else "").lower()
            if style.startswith("heading"):
                parts.append(f"\n# {txt}\n")
            else:
                parts.append(txt)
    for table_idx, table in enumerate(doc.tables, 1):
        parts.append(f"\n# Table {table_idx}\n")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return [("docx", "\n".join(parts))]


def read_pdf(path: Path) -> list[tuple[str, str]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError as exc:
            raise RuntimeError("pypdf or PyPDF2 is required to ingest .pdf files") from exc
    reader = PdfReader(str(path))
    pages: list[tuple[str, str]] = []
    for idx, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append((f"p.{idx}", text))
    return pages


def html_to_text(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</p\s*>", "\n", raw)
    raw = re.sub(r"(?i)</h[1-6]\s*>", "\n", raw)
    raw = re.sub(r"<[^>]+>", " ", raw)
    return normalize_ws(unescape(raw))


def read_epub(path: Path) -> list[tuple[str, str]]:
    parts: list[tuple[str, str]] = []
    with zipfile.ZipFile(path) as zf:
        container = ET.fromstring(zf.read("META-INF/container.xml"))
        rootfile = None
        for elem in container.iter():
            if elem.tag.endswith("rootfile"):
                rootfile = elem.attrib.get("full-path")
                break
        if not rootfile:
            raise RuntimeError("EPUB container does not include a rootfile")
        opf_dir = str(Path(rootfile).parent)
        opf_dir = "" if opf_dir == "." else opf_dir
        opf = ET.fromstring(zf.read(rootfile))
        manifest: dict[str, str] = {}
        spine: list[str] = []
        for elem in opf.iter():
            if elem.tag.endswith("item"):
                media = elem.attrib.get("media-type", "")
                if "html" in media or elem.attrib.get("href", "").lower().endswith((".xhtml", ".html", ".htm")):
                    manifest[elem.attrib.get("id", "")] = elem.attrib.get("href", "")
            elif elem.tag.endswith("itemref"):
                spine.append(elem.attrib.get("idref", ""))
        hrefs = [manifest[sid] for sid in spine if sid in manifest] or list(manifest.values())
        for idx, href in enumerate(hrefs, 1):
            member = str(Path(opf_dir) / href) if opf_dir else href
            try:
                raw = zf.read(member).decode("utf-8", errors="replace")
            except KeyError:
                continue
            text = html_to_text(raw)
            if text:
                parts.append((f"epub-section-{idx}", text))
    return parts


def read_json_material(path: Path) -> list[tuple[str, str]]:
    if path.suffix.lower() == ".jsonl":
        items = []
        for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
            line = line.strip()
            if line:
                items.append(json.loads(line))
    else:
        data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        items = data if isinstance(data, list) else [data]
    parts: list[tuple[str, str]] = []
    for idx, item in enumerate(items, 1):
        if not isinstance(item, dict):
            parts.append((f"item-{idx}", str(item)))
            continue
        title = item.get("title") or item.get("headline") or f"item-{idx}"
        date = item.get("date") or item.get("published_at") or item.get("publish_date")
        source = item.get("source") or item.get("source_name")
        url = item.get("url")
        text = item.get("full_text") or item.get("text") or item.get("content") or item.get("body") or ""
        entities = item.get("entities")
        tags = item.get("tags")
        summary = item.get("event_summary") or item.get("summary")
        structured = [
            f"# {title}",
            f"date: {date}" if date else "",
            f"source: {source}" if source else "",
            f"url: {url}" if url else "",
            f"entities: {json.dumps(entities, ensure_ascii=False)}" if entities else "",
            f"tags: {json.dumps(tags, ensure_ascii=False)}" if tags else "",
            f"event_summary: {summary}" if summary else "",
            str(text),
        ]
        parts.append((f"item-{idx}", "\n".join(x for x in structured if x)))
    return parts


def load_document(path: Path) -> list[tuple[str, str]]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return read_text_file(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".epub":
        return read_epub(path)
    if suffix in {".json", ".jsonl"}:
        return read_json_material(path)
    raise RuntimeError(f"unsupported file type: {path}")


HEADING_PATTERNS = [
    re.compile(r"^(#{1,6})\s+(.+)$"),
    re.compile(r"^第[一二三四五六七八九十百千万0-9]+[章节篇编部分卷]\s*[：:、.\-\s]*(.+)$"),
    re.compile(r"^[一二三四五六七八九十]+[、.]\s*(.+)$"),
    re.compile(r"^\d+(?:\.\d+){0,4}[、.\s]+(.+)$"),
]


def detect_heading(line: str) -> tuple[int, str] | None:
    stripped = line.strip()
    if not stripped or len(stripped) > 100:
        return None
    md = HEADING_PATTERNS[0].match(stripped)
    if md:
        return (len(md.group(1)), md.group(2).strip())
    for pattern in HEADING_PATTERNS[1:]:
        if pattern.match(stripped):
            level = 1
            if re.match(r"^\d+(?:\.\d+)+", stripped):
                level = stripped.split()[0].count(".") + 1 if " " in stripped else stripped.split("、")[0].count(".") + 1
            return (min(level, 6), stripped)
    return None


def sections_from_parts(parts: list[tuple[str, str]], title: str) -> tuple[list[Section], list[str]]:
    sections: list[Section] = []
    toc_titles: list[str] = []
    stack: list[tuple[int, str]] = []
    current_title = title
    current_level = 1
    current_lines: list[str] = []
    current_locator = parts[0][0] if parts else "text"

    def flush() -> None:
        nonlocal current_lines
        text = normalize_ws("\n".join(current_lines))
        if text:
            path = " / ".join(name for _, name in stack) or current_title
            sections.append(Section(current_title, path, text, current_locator, current_level))
        current_lines = []

    for locator, raw in parts:
        for line in raw.splitlines():
            heading = detect_heading(line)
            if heading:
                flush()
                level, heading_title = heading
                while stack and stack[-1][0] >= level:
                    stack.pop()
                stack.append((level, heading_title))
                toc_titles.append(" / ".join(name for _, name in stack))
                current_title = heading_title
                current_level = level
                current_locator = locator
            else:
                current_lines.append(line)
        if len(parts) > 1:
            current_lines.append(f"\n[locator: {locator}]\n")
    flush()
    if not sections:
        joined = "\n".join(raw for _, raw in parts)
        sections.append(Section(title, title, normalize_ws(joined), parts[0][0] if parts else "text", 1))
    return sections, toc_titles


def chunk_text(text: str, max_chars: int = 900, overlap_chars: int = 120) -> list[str]:
    paragraphs = split_paragraphs(text)
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 1 <= max_chars:
            current = f"{current}\n{para}".strip()
            continue
        if current:
            chunks.append(current)
        if len(para) > max_chars:
            start = 0
            while start < len(para):
                end = min(len(para), start + max_chars)
                chunks.append(para[start:end])
                start = max(end - overlap_chars, end)
        else:
            current = para
    if current:
        chunks.append(current)
    if overlap_chars and len(chunks) > 1:
        merged = []
        prev_tail = ""
        for chunk in chunks:
            if prev_tail:
                merged.append(f"{prev_tail}\n{chunk}")
            else:
                merged.append(chunk)
            prev_tail = chunk[-overlap_chars:]
        return merged
    return chunks


def collect_files(input_path: Path) -> list[Path]:
    if input_path.is_file():
        return [input_path]
    files = []
    for path in sorted(input_path.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES:
            files.append(path)
    return files


def insert_chunk(
    conn: sqlite3.Connection,
    *,
    chunk_id: str,
    source_id: str,
    parent_id: str | None,
    layer: str,
    path: str,
    title: str,
    locator: str,
    text: str,
    metadata: dict[str, Any] | None = None,
    embedding: list[float] | None = None,
) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO chunks
        (id, source_id, parent_id, layer, path, title, locator, text, token_count, embedding_json, metadata_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            chunk_id,
            source_id,
            parent_id,
            layer,
            path,
            title,
            locator,
            text,
            approx_tokens(text),
            json.dumps(embedding, ensure_ascii=False) if embedding is not None else None,
            json.dumps(metadata or {}, ensure_ascii=False),
            now_iso(),
        ),
    )
    conn.execute("DELETE FROM chunks_fts WHERE chunk_id = ?", (chunk_id,))
    conn.execute(
        "INSERT INTO chunks_fts(chunk_id, title, path, text) VALUES (?, ?, ?, ?)",
        (chunk_id, title, path, text),
    )


def batch_embed(texts: list[str], embed: bool, embed_url: str, embed_model: str, batch_size: int = 8) -> list[list[float] | None]:
    if not embed:
        return [None] * len(texts)
    vectors: list[list[float] | None] = []
    for start in range(0, len(texts), batch_size):
        batch = texts[start : start + batch_size]
        vectors.extend(llama_embed(batch, embed_url, embed_model))
    return vectors


def ingest_file(
    conn: sqlite3.Connection,
    path: Path,
    *,
    kind: str,
    title: str | None,
    source_name: str | None,
    url: str | None,
    published_at: str | None,
    version: str | None,
    embed: bool,
    embed_url: str,
    embed_model: str,
    max_chars: int,
) -> dict[str, Any]:
    doc_title = title or path.stem
    parts = load_document(path)
    sections, toc_titles = sections_from_parts(parts, doc_title)
    source_id = stable_id(str(path.resolve()), doc_title, kind, version or "")

    conn.execute("DELETE FROM sources WHERE id = ?", (source_id,))
    conn.execute(
        """
        INSERT INTO sources(id, kind, title, path, source_name, url, published_at, version, metadata_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            source_id,
            kind,
            doc_title,
            str(path.resolve()),
            source_name,
            url,
            published_at,
            version,
            json.dumps({"suffix": path.suffix.lower()}, ensure_ascii=False),
            now_iso(),
        ),
    )

    overview_text = "\n".join(section.text[:500] for section in sections[:3])[:1800]
    overview_id = stable_id(source_id, "overview")
    toc_text = "\n".join(toc_titles[:500]) or "\n".join(section.path for section in sections[:200])
    toc_id = stable_id(source_id, "toc")
    overview_vec, toc_vec = batch_embed([overview_text, toc_text], embed, embed_url, embed_model)
    insert_chunk(
        conn,
        chunk_id=overview_id,
        source_id=source_id,
        parent_id=None,
        layer="overview",
        path=doc_title,
        title=f"{doc_title} overview",
        locator="overview",
        text=overview_text or doc_title,
        embedding=overview_vec,
    )
    insert_chunk(
        conn,
        chunk_id=toc_id,
        source_id=source_id,
        parent_id=None,
        layer="toc",
        path=doc_title,
        title=f"{doc_title} TOC",
        locator="toc",
        text=toc_text or doc_title,
        embedding=toc_vec,
    )

    parent_count = 0
    child_count = 0
    for sidx, section in enumerate(sections, 1):
        parent_id = stable_id(source_id, "parent", str(sidx), section.path, section.text[:200])
        parent_vec = batch_embed([f"{section.path}\n{section.text[:1800]}"], embed, embed_url, embed_model)[0]
        insert_chunk(
            conn,
            chunk_id=parent_id,
            source_id=source_id,
            parent_id=None,
            layer="parent",
            path=section.path,
            title=section.title,
            locator=section.locator,
            text=section.text,
            metadata={"level": section.level},
            embedding=parent_vec,
        )
        parent_count += 1

        child_texts = chunk_text(section.text, max_chars=max_chars)
        child_vectors = batch_embed([f"{section.path}\n{txt}" for txt in child_texts], embed, embed_url, embed_model)
        for cidx, (child, child_vec) in enumerate(zip(child_texts, child_vectors), 1):
            child_id = stable_id(source_id, "content", str(sidx), str(cidx), section.path, child[:200])
            insert_chunk(
                conn,
                chunk_id=child_id,
                source_id=source_id,
                parent_id=parent_id,
                layer="content",
                path=section.path,
                title=section.title,
                locator=f"{section.locator}#chunk-{cidx}",
                text=child,
                metadata={"level": section.level, "chunk_index": cidx},
                embedding=child_vec,
            )
            child_count += 1
    conn.commit()
    return {"source_id": source_id, "title": doc_title, "parents": parent_count, "children": child_count}


def tokenize_query(query: str) -> list[str]:
    terms = re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z0-9_]{2,}", query)
    if not terms and query.strip():
        terms = [query.strip()]
    return terms[:12]


def fts_match_query(query: str) -> str:
    terms = tokenize_query(query)
    escaped = []
    for term in terms:
        term = term.replace('"', '""')
        escaped.append(f'"{term}"')
    return " OR ".join(escaped) if escaped else '""'


def cosine(a: list[float], b: list[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def load_vector(raw: str | None) -> list[float] | None:
    if not raw:
        return None
    try:
        return [float(x) for x in json.loads(raw)]
    except Exception:
        return None


def keyword_substring_score(query: str, text: str, path: str, title: str) -> float:
    haystack = f"{title}\n{path}\n{text}".lower()
    score = 0.0
    for term in tokenize_query(query):
        score += haystack.count(term.lower()) * 0.6
    if len(query.strip()) >= 2 and query.strip().lower() in haystack:
        score += 1.2
    return score


def retrieve_evidence(
    conn: sqlite3.Connection,
    query: str,
    *,
    top_k: int,
    embed_url: str | None,
    embed_model: str,
    layers: set[str] | None = None,
) -> list[Evidence]:
    layers = layers or {"overview", "toc", "parent", "content"}
    route_paths: set[str] = set()
    fts_q = fts_match_query(query)

    route_rows: list[sqlite3.Row] = []
    try:
        route_rows = conn.execute(
            """
            SELECT c.*, s.kind, s.title AS source_title, s.path AS source_path, s.source_name, s.url, s.published_at,
                   bm25(chunks_fts) AS bm25_rank
            FROM chunks_fts
            JOIN chunks c ON c.id = chunks_fts.chunk_id
            JOIN sources s ON s.id = c.source_id
            WHERE chunks_fts MATCH ? AND c.layer IN ('overview', 'toc')
            ORDER BY bm25_rank
            LIMIT 20
            """,
            (fts_q,),
        ).fetchall()
    except sqlite3.OperationalError:
        route_rows = []
    for row in route_rows:
        if row["path"]:
            route_paths.add(row["path"])
        for line in (row["text"] or "").splitlines():
            if any(term in line for term in tokenize_query(query)):
                route_paths.add(line.strip())

    candidates: dict[str, dict[str, Any]] = {}
    try:
        rows = conn.execute(
            """
            SELECT c.*, s.kind, s.title AS source_title, s.path AS source_path, s.source_name, s.url, s.published_at,
                   bm25(chunks_fts) AS bm25_rank
            FROM chunks_fts
            JOIN chunks c ON c.id = chunks_fts.chunk_id
            JOIN sources s ON s.id = c.source_id
            WHERE chunks_fts MATCH ?
            ORDER BY bm25_rank
            LIMIT 80
            """,
            (fts_q,),
        ).fetchall()
    except sqlite3.OperationalError:
        rows = []
    for row in rows:
        if row["layer"] not in layers:
            continue
        bm25_score = 1.0 / (1.0 + max(0.0, float(row["bm25_rank"])))
        candidates[row["id"]] = {"row": row, "score": bm25_score * 2.0}

    rows = conn.execute(
        """
        SELECT c.*, s.kind, s.title AS source_title, s.path AS source_path, s.source_name, s.url, s.published_at
        FROM chunks c
        JOIN sources s ON s.id = c.source_id
        WHERE c.layer IN ('parent', 'content')
        """
    ).fetchall()
    query_vec = None
    if embed_url:
        try:
            query_vec = llama_embed([query], embed_url, embed_model)[0]
        except Exception as exc:
            print(f"[warn] embedding retrieval disabled: {exc}", file=sys.stderr)
    for row in rows:
        if row["layer"] not in layers:
            continue
        score = keyword_substring_score(query, row["text"], row["path"] or "", row["title"] or "")
        if route_paths and any(rp and (rp in (row["path"] or "") or (row["path"] or "") in rp) for rp in route_paths):
            score += 0.8
        if query_vec:
            vec = load_vector(row["embedding_json"])
            if vec:
                score += max(0.0, cosine(query_vec, vec)) * 2.5
        if score <= 0:
            continue
        current = candidates.get(row["id"])
        if current:
            current["score"] += score
        else:
            candidates[row["id"]] = {"row": row, "score": score}

    expanded: dict[str, dict[str, Any]] = {}
    for item in candidates.values():
        row = item["row"]
        score = float(item["score"])
        if row["layer"] == "content" and row["parent_id"]:
            parent = conn.execute(
                """
                SELECT c.*, s.kind, s.title AS source_title, s.path AS source_path, s.source_name, s.url, s.published_at
                FROM chunks c
                JOIN sources s ON s.id = c.source_id
                WHERE c.id = ?
                """,
                (row["parent_id"],),
            ).fetchone()
            if parent:
                prev = expanded.get(parent["id"])
                if not prev or prev["score"] < score + 0.35:
                    expanded[parent["id"]] = {"row": parent, "score": score + 0.35}
        prev = expanded.get(row["id"])
        if not prev or prev["score"] < score:
            expanded[row["id"]] = {"row": row, "score": score}

    ranked = sorted(expanded.values(), key=lambda x: x["score"], reverse=True)
    evidences: list[Evidence] = []
    seen_text: set[str] = set()
    for idx, item in enumerate(ranked, 1):
        row = item["row"]
        digest = hashlib.sha1((row["text"] or "")[:500].encode("utf-8", errors="ignore")).hexdigest()
        if digest in seen_text:
            continue
        seen_text.add(digest)
        evidences.append(
            Evidence(
                id=f"E{len(evidences) + 1}",
                chunk_id=row["id"],
                source_id=row["source_id"],
                layer=row["layer"],
                path=row["path"] or "",
                title=row["title"] or "",
                locator=row["locator"] or "",
                text=row["text"] or "",
                score=round(float(item["score"]), 4),
                source_title=row["source_title"] or "",
                source_path=row["source_path"] or "",
                source_kind=row["kind"],
                source_name=row["source_name"],
                source_url=row["url"],
                published_at=row["published_at"],
            )
        )
        if len(evidences) >= top_k:
            break
    return evidences


def evidence_to_json(evidence: list[Evidence], max_chars: int = 1200) -> list[dict[str, Any]]:
    data = []
    for ev in evidence:
        data.append(
            {
                "id": ev.id,
                "chunk_id": ev.chunk_id,
                "layer": ev.layer,
                "score": ev.score,
                "source_kind": ev.source_kind,
                "role": "background_current_affairs" if ev.source_kind == "current_affairs" else "core_course_evidence",
                "citation": ev.citation(),
                "path": ev.path,
                "locator": ev.locator,
                "text": ev.text[:max_chars],
            }
        )
    return data


def has_locator(ev: Evidence) -> bool:
    return bool(ev.source_title or ev.source_path) and bool(ev.locator or ev.path)


def gate_evidence(evidence: list[Evidence], min_evidence: int, strict_current: bool = False) -> tuple[bool, dict[str, Any]]:
    final = [ev for ev in evidence if ev.layer in {"parent", "content"}]
    with_locators = [ev for ev in final if has_locator(ev)]
    issues = []
    if len(final) < min_evidence:
        issues.append(f"need at least {min_evidence} content/parent evidence chunks, got {len(final)}")
    if len(with_locators) < min_evidence:
        issues.append(f"need at least {min_evidence} cited evidence chunks, got {len(with_locators)}")
    if strict_current:
        dated = [ev for ev in with_locators if ev.published_at or ev.source_url]
        if len(dated) < min_evidence:
            issues.append("strict current-affairs mode requires dated URL/file-located sources")
    return not issues, {"ok": not issues, "issues": issues, "usable_evidence": len(with_locators)}


def build_generation_prompt(
    topic: str,
    question_type: str,
    count: int,
    difficulty: str,
    evidence: list[Evidence],
    language: str,
) -> list[dict[str, str]]:
    ev_json = evidence_to_json(evidence, max_chars=1600)
    schema = {
        "status": "ok",
        "topic": topic,
        "question_type": question_type,
        "items": [
            {
                "id": "Q1",
                "stem": "...",
                "material": "... optional",
                "options": [{"label": "A", "text": "...", "citations": ["E1"]}],
                "answer": "A",
                "analysis": "...",
                "citations": ["E1", "E2"],
                "assertions": [{"claim": "...", "citations": ["E1"]}],
                "difficulty": difficulty,
                "valid_until": None,
                "evidence_roles": {"core": ["E1"], "background": []},
            }
        ],
    }
    system = (
        "你是资深出题人和严格事实核验员。只能使用用户提供的 evidence JSON 出题。"
        "不得使用常识补充事实，不得编造日期、人名、机构、政策、页码、URL或引用。"
        "教材、讲义、书籍、课程大纲证据是 core_course_evidence；时政、热点、政策新闻素材是 background_current_affairs。"
        "除非用户明确要求纯时政题，否则题目考查点必须优先由 core_course_evidence 支撑，background_current_affairs 只能作为材料背景、案例或辅助解释。"
        "证据不足时输出 {\"status\":\"refused\",\"reason\":\"...\",\"missing_evidence\":[...]}。"
    )
    user = f"""
请基于 evidence JSON 生成考试题。

硬性规则：
1. 题干、答案、解析、材料和每个关键断言都必须引用 evidence id。
2. 错误选项必须有明确错因：被证据反驳、偷换概念、主体/时间/范围错误，或 evidence_not_supported。
3. 不要输出 evidence 中没有的事实。
4. current_affairs 证据必须保留来源、日期、URL或文件定位；若可能过时，为题目设置 valid_until 或在 analysis 中说明复检要求。
5. 每道题写 evidence_roles，区分 core 与 background；不要让 background_current_affairs 单独支撑教材知识点答案。
6. 只输出 JSON，不要 Markdown。
7. 语言：{language}。

参数：
- topic: {topic}
- question_type: {question_type}
- count: {count}
- difficulty: {difficulty}

必须匹配这个 JSON 形状：
{json.dumps(schema, ensure_ascii=False, indent=2)}

evidence:
{json.dumps(ev_json, ensure_ascii=False, indent=2)}
""".strip()
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def extract_json(text: str) -> Any:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start >= 0 and end > start:
            return json.loads(text[start : end + 1])
        raise


def verify_static(output: dict[str, Any], evidence: list[Evidence]) -> dict[str, Any]:
    evidence_ids = {ev.id for ev in evidence}
    issues: list[str] = []
    if output.get("status") == "refused":
        return {"ok": True, "mode": "static", "issues": [], "refused": True}
    if output.get("status") != "ok":
        issues.append("top-level status must be ok or refused")
    items = output.get("items")
    if not isinstance(items, list) or not items:
        issues.append("items must be a non-empty list")
        return {"ok": False, "mode": "static", "issues": issues}
    for i, item in enumerate(items, 1):
        prefix = f"item {i}"
        for field in ["stem", "answer", "analysis", "citations", "assertions"]:
            if field not in item:
                issues.append(f"{prefix}: missing {field}")
        citations = item.get("citations") or []
        if not isinstance(citations, list) or not citations:
            issues.append(f"{prefix}: citations must be a non-empty list")
        else:
            bad = [c for c in citations if c not in evidence_ids]
            if bad:
                issues.append(f"{prefix}: unknown citations {bad}")
        assertions = item.get("assertions") or []
        if not isinstance(assertions, list) or not assertions:
            issues.append(f"{prefix}: assertions must be a non-empty list")
        for aidx, assertion in enumerate(assertions, 1):
            ac = assertion.get("citations") if isinstance(assertion, dict) else None
            claim = assertion.get("claim") if isinstance(assertion, dict) else None
            if not claim:
                issues.append(f"{prefix} assertion {aidx}: missing claim")
            if not ac or not isinstance(ac, list):
                issues.append(f"{prefix} assertion {aidx}: missing citations")
            else:
                bad = [c for c in ac if c not in evidence_ids]
                if bad:
                    issues.append(f"{prefix} assertion {aidx}: unknown citations {bad}")
        if item.get("options"):
            labels = [opt.get("label") for opt in item["options"] if isinstance(opt, dict)]
            if item.get("answer") not in labels:
                issues.append(f"{prefix}: answer is not one of option labels")
    return {"ok": not issues, "mode": "static", "issues": issues}


def verify_with_llm(
    output: dict[str, Any],
    evidence: list[Evidence],
    llm_url: str,
    llm_model: str,
) -> dict[str, Any]:
    static = verify_static(output, evidence)
    if not static["ok"] or output.get("status") == "refused":
        return static
    messages = [
        {
            "role": "system",
            "content": (
                "你是严格证据核验器。只判断题目 JSON 中每个断言是否被 evidence 支持。"
                "输出 JSON: {\"ok\":true/false,\"issues\":[...],\"item_results\":[...]}"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "task": "verify_exam_items_against_evidence",
                    "evidence": evidence_to_json(evidence, max_chars=1800),
                    "output": output,
                    "rules": [
                        "unsupported or contradicted assertions are failures",
                        "missing citations are failures",
                        "answer key inconsistent with evidence is a failure",
                        "do not repair; only verify",
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
        },
    ]
    raw = llama_chat(messages, llm_url, llm_model, temperature=0.0, max_tokens=2048)
    try:
        llm_report = extract_json(raw)
    except Exception:
        return {"ok": False, "mode": "llm", "issues": ["verifier returned non-JSON", raw[:800]], "static": static}
    if not isinstance(llm_report, dict):
        return {"ok": False, "mode": "llm", "issues": ["verifier returned non-object JSON"], "static": static}
    llm_report["static"] = static
    llm_report["mode"] = "llm"
    llm_report["ok"] = bool(llm_report.get("ok")) and static["ok"]
    return llm_report


def refusal(topic: str, gate: dict[str, Any], evidence: list[Evidence]) -> dict[str, Any]:
    return {
        "status": "refused",
        "topic": topic,
        "reason": "evidence_gate_failed",
        "missing_evidence": gate.get("issues", []),
        "strongest_evidence": evidence_to_json(evidence[:5], max_chars=700),
    }


def rewrite_once(
    output: dict[str, Any],
    verification: dict[str, Any],
    topic: str,
    question_type: str,
    count: int,
    difficulty: str,
    evidence: list[Evidence],
    language: str,
    llm_url: str,
    llm_model: str,
) -> dict[str, Any]:
    messages = build_generation_prompt(topic, question_type, count, difficulty, evidence, language)
    messages.append(
        {
            "role": "user",
            "content": (
                "上一次输出未通过核验。请只使用同一份 evidence 重写，修复所有问题。"
                "如果无法修复，输出 refused JSON。\n\n"
                f"verification:\n{json.dumps(verification, ensure_ascii=False, indent=2)}\n\n"
                f"previous_output:\n{json.dumps(output, ensure_ascii=False, indent=2)}"
            ),
        }
    )
    raw = llama_chat(messages, llm_url, llm_model, temperature=0.1, max_tokens=4096)
    parsed = extract_json(raw)
    if not isinstance(parsed, dict):
        raise RuntimeError("rewrite returned non-object JSON")
    return parsed


def cmd_init_db(args: argparse.Namespace) -> None:
    conn = connect(args.db)
    init_db(conn)
    print(json.dumps({"ok": True, "db": args.db}, ensure_ascii=False, indent=2))


def cmd_ingest(args: argparse.Namespace) -> None:
    conn = connect(args.db)
    ensure_db(conn)
    files = collect_files(Path(args.input))
    if not files:
        raise SystemExit(f"no supported files found in {args.input}")
    results = []
    for path in files:
        print(f"[ingest] {path}", file=sys.stderr)
        result = ingest_file(
            conn,
            path,
            kind=args.kind,
            title=args.title if len(files) == 1 else None,
            source_name=args.source_name,
            url=args.url,
            published_at=args.published_at,
            version=args.version,
            embed=args.embed,
            embed_url=args.embed_url,
            embed_model=args.embed_model,
            max_chars=args.max_chars,
        )
        results.append(result)
    print(json.dumps({"ok": True, "files": len(files), "results": results}, ensure_ascii=False, indent=2))


def cmd_retrieve(args: argparse.Namespace) -> None:
    conn = connect(args.db)
    ensure_db(conn)
    embed_url = args.embed_url if args.embed_url else None
    evidence = retrieve_evidence(conn, args.query, top_k=args.top_k, embed_url=embed_url, embed_model=args.embed_model)
    print(json.dumps({"query": args.query, "evidence": evidence_to_json(evidence, max_chars=args.max_chars)}, ensure_ascii=False, indent=2))


def cmd_generate(args: argparse.Namespace) -> None:
    conn = connect(args.db)
    ensure_db(conn)
    evidence = retrieve_evidence(
        conn,
        args.topic,
        top_k=args.top_k,
        embed_url=args.embed_url,
        embed_model=args.embed_model,
        layers={"parent", "content", "overview", "toc"},
    )
    ok, gate = gate_evidence(evidence, args.min_evidence, args.strict_current)
    prompt_record = {
        "topic": args.topic,
        "question_type": args.question_type,
        "count": args.count,
        "difficulty": args.difficulty,
        "min_evidence": args.min_evidence,
        "strict_current": args.strict_current,
    }
    if not ok:
        output = refusal(args.topic, gate, evidence)
        verification = {"ok": False, "mode": "gate", "issues": gate["issues"]}
        status = "refused"
    else:
        messages = build_generation_prompt(args.topic, args.question_type, args.count, args.difficulty, evidence, args.language)
        raw = llama_chat(messages, args.llm_url, args.llm_model, temperature=args.temperature, max_tokens=args.max_tokens)
        try:
            output = extract_json(raw)
        except Exception as exc:
            output = {"status": "refused", "reason": "generation_returned_non_json", "raw": raw[:2000], "error": str(exc)}
            verification = {"ok": False, "mode": "parse", "issues": ["generation returned non-JSON"]}
            status = "refused"
        else:
            if not isinstance(output, dict):
                output = {"status": "refused", "reason": "generation_returned_non_object_json", "raw": output}
            verification = (
                verify_with_llm(output, evidence, args.llm_url, args.llm_model)
                if args.llm_verify
                else verify_static(output, evidence)
            )
            if not verification.get("ok") and args.rewrite_on_fail and output.get("status") != "refused":
                try:
                    output = rewrite_once(
                        output,
                        verification,
                        args.topic,
                        args.question_type,
                        args.count,
                        args.difficulty,
                        evidence,
                        args.language,
                        args.llm_url,
                        args.llm_model,
                    )
                    verification = (
                        verify_with_llm(output, evidence, args.llm_url, args.llm_model)
                        if args.llm_verify
                        else verify_static(output, evidence)
                    )
                except Exception as exc:
                    verification = {"ok": False, "mode": "rewrite", "issues": [f"rewrite failed: {exc}"]}
            status = "ok" if verification.get("ok") and output.get("status") == "ok" else "refused"
            if status != "ok" and output.get("status") != "refused":
                output = {
                    "status": "refused",
                    "reason": "verification_failed",
                    "verification": verification,
                    "strongest_evidence": evidence_to_json(evidence[:5], max_chars=700),
                    "draft": output,
                }

    qid = str(uuid.uuid4())
    conn.execute(
        """
        INSERT INTO questions(id, topic, question_type, prompt_json, evidence_json, output_json, verification_json, status, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            qid,
            args.topic,
            args.question_type,
            json.dumps(prompt_record, ensure_ascii=False),
            json.dumps(evidence_to_json(evidence), ensure_ascii=False),
            json.dumps(output, ensure_ascii=False),
            json.dumps(verification, ensure_ascii=False),
            status,
            now_iso(),
        ),
    )
    conn.commit()
    result = {
        "id": qid,
        "status": status,
        "gate": gate,
        "verification": verification,
        "evidence": evidence_to_json(evidence, max_chars=900),
        "output": output,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


def cmd_stats(args: argparse.Namespace) -> None:
    conn = connect(args.db)
    ensure_db(conn)
    data = {
        "sources": conn.execute("SELECT COUNT(*) FROM sources").fetchone()[0],
        "chunks": conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0],
        "questions": conn.execute("SELECT COUNT(*) FROM questions").fetchone()[0],
        "layers": {
            row["layer"]: row["n"]
            for row in conn.execute("SELECT layer, COUNT(*) AS n FROM chunks GROUP BY layer").fetchall()
        },
    }
    print(json.dumps(data, ensure_ascii=False, indent=2))


def cmd_vacuum(args: argparse.Namespace) -> None:
    conn = connect(args.db)
    conn.execute("VACUUM")
    print(json.dumps({"ok": True, "db": args.db, "action": "vacuum"}, ensure_ascii=False, indent=2))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Local evidence-gated senior exam writer with SQLite and llama.cpp.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent(
            """
            Examples:
              python scripts/senior_exam_writer.py init-db --db ./exam.sqlite
              python scripts/senior_exam_writer.py ingest --db ./exam.sqlite --input ./materials --kind book --embed
              python scripts/senior_exam_writer.py retrieve --db ./exam.sqlite --query "共同富裕"
              python scripts/senior_exam_writer.py generate --db ./exam.sqlite --topic "共同富裕" --count 3 --llm-verify
            """
        ),
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("init-db", help="create or migrate the SQLite evidence database")
    p.add_argument("--db", required=True)
    p.set_defaults(func=cmd_init_db)

    p = sub.add_parser("ingest", help="ingest files into the evidence database")
    p.add_argument("--db", required=True)
    p.add_argument("--input", required=True, help="file or directory")
    p.add_argument("--kind", default="book", choices=["book", "handout", "outline", "current_affairs", "notes"])
    p.add_argument("--title")
    p.add_argument("--source-name")
    p.add_argument("--url")
    p.add_argument("--published-at")
    p.add_argument("--version")
    p.add_argument("--embed", action="store_true", help="call local llama.cpp embedding server")
    p.add_argument("--no-embed", action="store_false", dest="embed")
    p.add_argument("--embed-url", default=DEFAULT_EMBED_URL)
    p.add_argument("--embed-model", default=DEFAULT_EMBED_MODEL)
    p.add_argument("--max-chars", type=int, default=900)
    p.set_defaults(func=cmd_ingest, embed=False)

    p = sub.add_parser("retrieve", help="retrieve evidence for inspection")
    p.add_argument("--db", required=True)
    p.add_argument("--query", required=True)
    p.add_argument("--top-k", type=int, default=8)
    p.add_argument("--embed-url", default=DEFAULT_EMBED_URL, help="set empty string to disable semantic query embedding")
    p.add_argument("--embed-model", default=DEFAULT_EMBED_MODEL)
    p.add_argument("--max-chars", type=int, default=900)
    p.set_defaults(func=cmd_retrieve)

    p = sub.add_parser("generate", help="generate evidence-gated exam questions")
    p.add_argument("--db", required=True)
    p.add_argument("--topic", required=True)
    p.add_argument("--question-type", default="single_choice", choices=["single_choice", "multiple_choice", "material_analysis", "short_answer"])
    p.add_argument("--count", type=int, default=3)
    p.add_argument("--difficulty", default="medium", choices=["easy", "medium", "hard"])
    p.add_argument("--language", default="zh-CN")
    p.add_argument("--top-k", type=int, default=10)
    p.add_argument("--min-evidence", type=int, default=3)
    p.add_argument("--strict-current", action="store_true")
    p.add_argument("--embed-url", default=DEFAULT_EMBED_URL)
    p.add_argument("--embed-model", default=DEFAULT_EMBED_MODEL)
    p.add_argument("--llm-url", default=DEFAULT_LLM_URL)
    p.add_argument("--llm-model", default=DEFAULT_LLM_MODEL)
    p.add_argument("--temperature", type=float, default=0.2)
    p.add_argument("--max-tokens", type=int, default=4096)
    p.add_argument("--llm-verify", action="store_true")
    p.add_argument("--rewrite-on-fail", action="store_true", default=True)
    p.add_argument("--no-rewrite-on-fail", action="store_false", dest="rewrite_on_fail")
    p.set_defaults(func=cmd_generate)

    p = sub.add_parser("stats", help="show database counts")
    p.add_argument("--db", required=True)
    p.set_defaults(func=cmd_stats)

    p = sub.add_parser("vacuum", help="vacuum the SQLite database")
    p.add_argument("--db", required=True)
    p.set_defaults(func=cmd_vacuum)
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    started = time.time()
    try:
        args.func(args)
    except KeyboardInterrupt:
        print("interrupted", file=sys.stderr)
        return 130
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False, indent=2), file=sys.stderr)
        return 1
    finally:
        elapsed = time.time() - started
        if elapsed > 5:
            print(f"[done] {elapsed:.1f}s", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
