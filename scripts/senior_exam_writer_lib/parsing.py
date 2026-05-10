from __future__ import annotations

import json
import re
import zipfile
from html import unescape
from pathlib import Path
from xml.etree import ElementTree as ET

from .common import Section

def approx_tokens(text: str) -> int:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    words = len(re.findall(r"[A-Za-z0-9_]+", text))
    punctuation = len(re.findall(r"[^\w\s\u4e00-\u9fff]", text))
    return cjk + words + max(1, punctuation // 3)

def normalize_ws(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)

def split_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"\n{1,}", normalize_ws(text)) if p.strip()]

def read_text_file(path: Path) -> list[tuple[str, str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [("text", text)]

def read_docx(path: Path) -> list[tuple[str, str]]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required to ingest .docx files") from exc
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            style = (para.style.name if para.style else "").lower()
            if style.startswith("heading"):
                parts.append(f"\n# {txt}\n")
            else:
                parts.append(txt)
    for table_idx, table in enumerate(doc.tables, 1):
        parts.append(f"\n# Table {table_idx}\n")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return [("docx", "\n".join(parts))]

def read_pdf(path: Path) -> list[tuple[str, str]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError as exc:
            raise RuntimeError("pypdf or PyPDF2 is required to ingest .pdf files") from exc
    reader = PdfReader(str(path))
    pages: list[tuple[str, str]] = []
    for idx, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append((f"p.{idx}", text))
    return pages

def html_to_text(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</p\s*>", "\n", raw)
    raw = re.sub(r"(?i)</h[1-6]\s*>", "\n", raw)
    raw = re.sub(r"<[^>]+>", " ", raw)
    return normalize_ws(unescape(raw))

def read_epub(path: Path) -> list[tuple[str, str]]:
    parts: list[tuple[str, str]] = []
    with zipfile.ZipFile(path) as zf:
        container = ET.fromstring(zf.read("META-INF/container.xml"))
        rootfile = None
        for elem in container.iter():
            if elem.tag.endswith("rootfile"):
                rootfile = elem.attrib.get("full-path")
                break
        if not rootfile:
            raise RuntimeError("EPUB container does not include a rootfile")
        opf_dir = str(Path(rootfile).parent)
        opf_dir = "" if opf_dir == "." else opf_dir
        opf = ET.fromstring(zf.read(rootfile))
        manifest: dict[str, str] = {}
        spine: list[str] = []
        for elem in opf.iter():
            if elem.tag.endswith("item"):
                media = elem.attrib.get("media-type", "")
                if "html" in media or elem.attrib.get("href", "").lower().endswith((".xhtml", ".html", ".htm")):
                    manifest[elem.attrib.get("id", "")] = elem.attrib.get("href", "")
            elif elem.tag.endswith("itemref"):
                spine.append(elem.attrib.get("idref", ""))
        hrefs = [manifest[sid] for sid in spine if sid in manifest] or list(manifest.values())
        for idx, href in enumerate(hrefs, 1):
            member = str(Path(opf_dir) / href) if opf_dir else href
            try:
                raw = zf.read(member).decode("utf-8", errors="replace")
            except KeyError:
                continue
            text = html_to_text(raw)
            if text:
                parts.append((f"epub-section-{idx}", text))
    return parts

def read_json_material(path: Path) -> list[tuple[str, str]]:
    if path.suffix.lower() == ".jsonl":
        items = []
        for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
            line = line.strip()
            if line:
                items.append(json.loads(line))
    else:
        data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        items = data if isinstance(data, list) else [data]
    parts: list[tuple[str, str]] = []
    for idx, item in enumerate(items, 1):
        if not isinstance(item, dict):
            parts.append((f"item-{idx}", str(item)))
            continue
        title = item.get("title") or item.get("headline") or f"item-{idx}"
        date = item.get("date") or item.get("published_at") or item.get("publish_date")
        source = item.get("source") or item.get("source_name")
        url = item.get("url")
        text = item.get("full_text") or item.get("text") or item.get("content") or item.get("body") or ""
        entities = item.get("entities")
        tags = item.get("tags")
        summary = item.get("event_summary") or item.get("summary")
        structured = [
            f"# {title}",
            f"date: {date}" if date else "",
            f"source: {source}" if source else "",
            f"url: {url}" if url else "",
            f"entities: {json.dumps(entities, ensure_ascii=False)}" if entities else "",
            f"tags: {json.dumps(tags, ensure_ascii=False)}" if tags else "",
            f"event_summary: {summary}" if summary else "",
            str(text),
        ]
        parts.append((f"item-{idx}", "\n".join(x for x in structured if x)))
    return parts

def load_document(path: Path) -> list[tuple[str, str]]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return read_text_file(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".epub":
        return read_epub(path)
    if suffix in {".json", ".jsonl"}:
        return read_json_material(path)
    raise RuntimeError(f"unsupported file type: {path}")

HEADING_PATTERNS = [
    re.compile(r"^(#{1,6})\s+(.+)$"),
    re.compile(r"^第[一二三四五六七八九十百千万0-9]+[章节篇编部分卷]\s*[：:、.\-\s]*(.+)$"),
    re.compile(r"^[一二三四五六七八九十]+[、.]\s*(.+)$"),
    re.compile(r"^\d+(?:\.\d+){0,4}[、.\s]+(.+)$"),
]

def detect_heading(line: str) -> tuple[int, str] | None:
    stripped = line.strip()
    if not stripped or len(stripped) > 100:
        return None
    md = HEADING_PATTERNS[0].match(stripped)
    if md:
        return (len(md.group(1)), md.group(2).strip())
    for pattern in HEADING_PATTERNS[1:]:
        if pattern.match(stripped):
            level = 1
            if re.match(r"^\d+(?:\.\d+)+", stripped):
                level = stripped.split()[0].count(".") + 1 if " " in stripped else stripped.split("、")[0].count(".") + 1
            return (min(level, 6), stripped)
    return None

def sections_from_parts(parts: list[tuple[str, str]], title: str) -> tuple[list[Section], list[str]]:
    sections: list[Section] = []
    toc_titles: list[str] = []
    stack: list[tuple[int, str]] = []
    current_title = title
    current_level = 1
    current_lines: list[str] = []
    current_locator = parts[0][0] if parts else "text"

    def flush() -> None:
        nonlocal current_lines
        text = normalize_ws("\n".join(current_lines))
        if text:
            path = " / ".join(name for _, name in stack) or current_title
            sections.append(Section(current_title, path, text, current_locator, current_level))
        current_lines = []

    for locator, raw in parts:
        for line in raw.splitlines():
            heading = detect_heading(line)
            if heading:
                flush()
                level, heading_title = heading
                while stack and stack[-1][0] >= level:
                    stack.pop()
                stack.append((level, heading_title))
                toc_titles.append(" / ".join(name for _, name in stack))
                current_title = heading_title
                current_level = level
                current_locator = locator
            else:
                current_lines.append(line)
        if len(parts) > 1:
            current_lines.append(f"\n[locator: {locator}]\n")
    flush()
    if not sections:
        joined = "\n".join(raw for _, raw in parts)
        sections.append(Section(title, title, normalize_ws(joined), parts[0][0] if parts else "text", 1))
    return sections, toc_titles

def chunk_text(text: str, max_chars: int = 900, overlap_chars: int = 120) -> list[str]:
    paragraphs = split_paragraphs(text)
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 1 <= max_chars:
            current = f"{current}\n{para}".strip()
            continue
        if current:
            chunks.append(current)
        if len(para) > max_chars:
            start = 0
            while start < len(para):
                end = min(len(para), start + max_chars)
                chunks.append(para[start:end])
                start = max(end - overlap_chars, end)
        else:
            current = para
    if current:
        chunks.append(current)
    if overlap_chars and len(chunks) > 1:
        merged = []
        prev_tail = ""
        for chunk in chunks:
            if prev_tail:
                merged.append(f"{prev_tail}\n{chunk}")
            else:
                merged.append(chunk)
            prev_tail = chunk[-overlap_chars:]
        return merged
    return chunks

